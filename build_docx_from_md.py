from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(".md/20260630_152848_result.md")
OUTPUT_DIR = Path(".docx")
OUTPUT = OUTPUT_DIR / "20260630_152848_result.docx"


def set_spacing(style, before=0, after=6, line=1.10):
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def set_run_defaults(style, font_name="Calibri", size=11, color="000000", bold=False):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.color.rgb = RGBColor.from_string(color)
    font.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = 2
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def parse_markdown(text):
    blocks = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        else:
            blocks.append(("paragraph", line))
    return blocks


def build_docx():
    source_text = SOURCE.read_text(encoding="utf-8-sig")
    OUTPUT_DIR.mkdir(exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_run_defaults(styles["Normal"], size=11)
    set_spacing(styles["Normal"], after=6, line=1.10)

    set_run_defaults(styles["Title"], size=18, color="0B2545", bold=True)
    set_spacing(styles["Title"], before=0, after=10, line=1.10)

    set_run_defaults(styles["Heading 1"], size=16, color="2E74B5", bold=True)
    set_spacing(styles["Heading 1"], before=16, after=8, line=1.10)

    set_run_defaults(styles["List Bullet"], size=11)
    set_spacing(styles["List Bullet"], after=6, line=1.10)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    for kind, content in parse_markdown(source_text):
        if kind == "title":
            p = doc.add_paragraph(style="Title")
            p.add_run(content)
        elif kind == "h1":
            doc.add_heading(content, level=1)
        elif kind == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        else:
            doc.add_paragraph(content)

    doc.core_properties.title = "2026-06-30 生活・研究ログ整理"
    doc.core_properties.subject = "生活・研究ログ"
    doc.core_properties.author = "OpenAI-Agent"
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_docx()
